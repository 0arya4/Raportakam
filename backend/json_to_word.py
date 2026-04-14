"""
Convert JSON report structure to Microsoft Word document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import io
KURDISH_FONT_PRIMARY = "Rabar_022"
KURDISH_FONT_FALLBACK = "NRT"


def _is_rtl_text(text):
    """Check if text contains RTL characters (Arabic/Kurdish)"""
    for ch in text[:100]:
        if '\u0600' <= ch <= '\u06FF' or '\u0750' <= ch <= '\u077F' or '\uFB50' <= ch <= '\uFDFF':
            return True
    return False


def _set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def _get_font_with_fallback(font_name):
    """Get font with fallback chain for RTL/Kurdish fonts"""
    # Fallback chain for Kurdish fonts
    if font_name == KURDISH_FONT_PRIMARY:
        return font_name, KURDISH_FONT_FALLBACK  # (primary, fallback)
    elif font_name == KURDISH_FONT_FALLBACK:
        return font_name, "Arial"  # Fallback to Arial if neither Kurdish font available
    return font_name, None


def _set_run_font(run, font_name, size_pt=None):
    """Set font for a run including complex script (RTL) font with fallback"""
    primary_font, fallback_font = _get_font_with_fallback(font_name)

    run.font.name = primary_font
    if fallback_font:
        run.font.name = f"{primary_font}, {fallback_font}"

    if size_pt:
        run.font.size = Pt(size_pt)

    # Set complex script font for RTL with fallback
    rPr = run._element.get_or_add_rPr()
    cs_font = OxmlElement('w:rFonts')
    cs_font.set(qn('w:cs'), primary_font)
    cs_font.set(qn('w:ascii'), primary_font)
    cs_font.set(qn('w:hAnsi'), primary_font)
    if fallback_font:
        cs_font.set(qn('w:eastAsia'), fallback_font)
    rPr.append(cs_font)


def _set_section_rtl(section):
    """Set document section to RTL"""
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    sectPr.append(bidi)


def _apply_font_to_paragraph(paragraph, font_name, size_pt=None):
    """Apply font to all runs in a paragraph"""
    for run in paragraph.runs:
        _set_run_font(run, font_name, size_pt)


def _strip_theme_color(run):
    """Remove theme color attributes from a run so RGB black is honored"""
    rPr = run._element.get_or_add_rPr()
    # Remove <w:color> elements that carry themeColor attribute
    for color_elem in rPr.findall(qn('w:color')):
        if color_elem.get(qn('w:themeColor')):
            color_elem.attrib.pop(qn('w:themeColor'), None)
            color_elem.attrib.pop(qn('w:themeShade'), None)
            color_elem.attrib.pop(qn('w:themeTint'), None)
        color_elem.set(qn('w:val'), '000000')


def _force_heading_black(doc):
    """Force all heading styles to be black color to avoid default blue"""
    from docx.shared import RGBColor
    black = RGBColor(0, 0, 0)
    for i in range(10):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.color.rgb = black
            # Strip theme color from the style XML itself
            try:
                rPr = style._element.get_or_add_rPr()
                for color_elem in rPr.findall(qn('w:color')):
                    color_elem.attrib.pop(qn('w:themeColor'), None)
                    color_elem.attrib.pop(qn('w:themeShade'), None)
                    color_elem.attrib.pop(qn('w:themeTint'), None)
                    color_elem.set(qn('w:val'), '000000')
            except Exception:
                pass
            # Remove any blue borders/lines often associated with headings
            try:
                pPr = style._element.get_or_add_pPr()
                # Remove existing borders first
                for existing_bdr in pPr.findall(qn('w:pBdr')):
                    pPr.remove(existing_bdr)
                pbdr = OxmlElement('w:pBdr')
                for side in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    pbdr.append(border)
                pPr.append(pbdr)
            except Exception:
                pass
    
    if 'Title' in doc.styles:
        doc.styles['Title'].font.color.rgb = black
        try:
            rPr = doc.styles['Title']._element.get_or_add_rPr()
            for color_elem in rPr.findall(qn('w:color')):
                color_elem.attrib.pop(qn('w:themeColor'), None)
                color_elem.set(qn('w:val'), '000000')
        except Exception:
            pass
    
    # Also force normal and list styles
    for style_name in ['Normal', 'Body Text', 'List Number', 'List Bullet', 'List Number 2']:
        if style_name in doc.styles:
            doc.styles[style_name].font.color.rgb = black

def _add_styled_heading(doc, text, level, font_name=None, is_rtl=False, page_break=False):
    """Add heading with black color and optional font/RTL"""
    if page_break:
        doc.add_page_break()

    # Strip [Heading X] labels
    if text.startswith("[Heading"):
        parts = text.split("]", 1)
        if len(parts) > 1:
            text = parts[1].strip()

    heading = doc.add_heading(text, level=level)

    # Set proper spacing and font for headings
    pf = heading.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        _strip_theme_color(run)
        # Use Calibri for English headings
        if not font_name and not is_rtl:
            run.font.name = "Calibri"
        elif font_name:
            _set_run_font(run, font_name)

    if is_rtl:
        _set_rtl_paragraph(heading)
    return heading


def _add_styled_paragraph(doc, text, font_name=None, size_pt=None, is_rtl=False, alignment=None, style=None):
    """Add paragraph with optional font/RTL/alignment"""
    p = doc.add_paragraph(text, style=style)

    # Apply formatting
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # Default to justified unless center/right aligned
    if alignment:
        p.alignment = alignment
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if is_rtl:
        _set_rtl_paragraph(p)

    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0) # Force black
        _strip_theme_color(run)
        if font_name:
            _set_run_font(run, font_name, size_pt)
        else:
            # Default to Times New Roman
            run.font.name = "Times New Roman"
            if size_pt:
                run.font.size = Pt(size_pt)
    return p


def add_table_to_doc(doc, table_data, font_name=None, is_rtl=False):
    """Add a formatted table to the document"""
    if not table_data or not table_data.get("headers"):
        return

    headers = table_data["headers"]
    rows = table_data.get("rows", [])

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    # Set all borders to black
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Add headers with black background and bold text
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                if font_name:
                    _set_run_font(run, font_name)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if is_rtl:
                _set_rtl_paragraph(paragraph)
        # Light gray background for header (with black text)
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "D3D3D3")  # Light gray instead of black
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)

    # Add data rows with black text, no alternating colors
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            # Ensure black text in all cells
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                    if font_name:
                        _set_run_font(run, font_name)
                if is_rtl:
                    _set_rtl_paragraph(paragraph)


def add_smart_structure_to_doc(doc, structure_data, font_name=None, is_rtl=False):
    """Add SmartArt-style structure (hierarchy or process)"""
    if not structure_data:
        return

    title = structure_data.get("title", "Structure")
    struct_type = structure_data.get("type", "hierarchy")
    data = structure_data.get("data", {})

    _add_styled_heading(doc, title, level=3, font_name=font_name, is_rtl=is_rtl)

    if struct_type == "process":
        steps = data.get("steps", [])
        for step in steps:
            step_num = step.get("step", "")
            step_name = step.get("name", "")
            step_desc = step.get("description", "")
            _add_styled_paragraph(doc, f"→ Step {step_num}: {step_name}", font_name=font_name, is_rtl=is_rtl, style="List Bullet")
            if step_desc:
                _add_styled_paragraph(doc, step_desc, font_name=font_name, is_rtl=is_rtl, style="List Bullet 2")

    elif struct_type == "hierarchy":
        root = data.get("root", "Root")
        children = data.get("children", [])

        p = _add_styled_paragraph(doc, root, font_name=font_name, is_rtl=is_rtl)
        p.paragraph_format.left_indent = Inches(0)

        def add_tree_item(parent_dict, indent_level=1):
            children_list = parent_dict.get("children", [])
            for child in children_list:
                p = _add_styled_paragraph(doc, f"└─ {child.get('name', '')}", font_name=font_name, is_rtl=is_rtl)
                p.paragraph_format.left_indent = Inches(indent_level * 0.3)
                if "children" in child:
                    add_tree_item(child, indent_level + 1)

        for child in children:
            p = _add_styled_paragraph(doc, f"├─ {child.get('name', '')}", font_name=font_name, is_rtl=is_rtl)
            p.paragraph_format.left_indent = Inches(0.3)
            if "children" in child:
                add_tree_item(child, indent_level=2)


def add_chart_description_to_doc(doc, chart_data, font_name=None, is_rtl=False):
    """Add chart as text representation"""
    if not chart_data:
        return

    title = chart_data.get("title", "Chart")
    chart_type = chart_data.get("type", "bar")
    labels = chart_data.get("labels", [])
    values = chart_data.get("values", [])
    description = chart_data.get("description", "")

    p = _add_styled_paragraph(doc, f"[{chart_type.upper()} CHART] {title}", font_name=font_name, is_rtl=is_rtl)
    p.runs[0].font.bold = True

    if labels and values:
        data_table = doc.add_table(rows=len(labels) + 1, cols=2)
        data_table.style = "Table Grid"

        # Set all borders to black
        tbl = data_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Add table borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        for cell_idx, header_text in enumerate(["Category", "Value"]):
            data_table.rows[0].cells[cell_idx].text = header_text
            for paragraph in data_table.rows[0].cells[cell_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                    if font_name:
                        _set_run_font(run, font_name)
            # Light gray background for header
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "D3D3D3")  # Light gray
            data_table.rows[0].cells[cell_idx]._element.get_or_add_tcPr().append(shading_elm)

        for i, (label, value) in enumerate(zip(labels, values)):
            data_table.rows[i + 1].cells[0].text = str(label)
            data_table.rows[i + 1].cells[1].text = str(value)
            for cell_idx in range(2):
                for paragraph in data_table.rows[i + 1].cells[cell_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                        if font_name:
                            _set_run_font(run, font_name)

    if description:
        _add_styled_paragraph(doc, f"Description: {description}", font_name=font_name, is_rtl=is_rtl)


def json_to_word(report_json_str, language="English"):
    """
    Convert JSON report to Word document

    Args:
        report_json_str: JSON string containing report structure
        language: Language for font selection ("English", "Kurdish (Sorani)", "Arabic")

    Returns:
        BytesIO object containing Word document
    """
    try:
        report_data = json.loads(report_json_str)
    except json.JSONDecodeError as e:
        print(f"[JSON] Decode error: {e}")
        print(f"[JSON] String length: {len(report_json_str)}")
        print(f"[JSON] First 200: {report_json_str[:200]}")
        print(f"[JSON] Last 100: {report_json_str[-100:]}")

        report_data = None
        if len(report_json_str) > 100:
            lines = report_json_str.split('\n')
            for trim_count in range(1, min(len(lines), 50)):
                candidate = '\n'.join(lines[:len(lines) - trim_count]).rstrip().rstrip(',')
                open_braces = candidate.count('{') - candidate.count('}')
                open_brackets = candidate.count('[') - candidate.count(']')
                closer = ']' * max(0, open_brackets) + '}' * max(0, open_braces)
                attempt = candidate + closer
                try:
                    report_data = json.loads(attempt)
                    print(f"[JSON] Recovered by trimming {trim_count} lines and closing {len(closer)} brackets")
                    break
                except json.JSONDecodeError:
                    continue
        if report_data is None:
            raise ValueError(f"Invalid JSON format: {str(e)} (could not recover)")

    # Detect language and set fonts
    is_rtl = language in ["Kurdish (Sorani)", "Arabic"]
    if not is_rtl:
        # Auto-detect from title
        title_text = report_data.get("title", "")
        is_rtl = _is_rtl_text(title_text)

    if is_rtl:
        # Use consistent Kurdish font with fallback
        body_font = KURDISH_FONT_PRIMARY  # "Rabar_022"
        heading_font = body_font
        body_size = 14
    else:
        body_font = "Times New Roman"
        heading_font = None  # Use Calibri for English headings
        body_size = 12

    # Create document
    doc = Document()
    _force_heading_black(doc)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if is_rtl:
            _set_section_rtl(section)

    # === COVER PAGE ===
    cover_data = report_data.get("cover", {})
    title = report_data.get("title", "Academic Report")

    # Add spacing at top
    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = title_para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(24)
    pf.line_spacing = 1.5
    for run in title_para.runs:
        run.font.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if not heading_font and not is_rtl:
            run.font.name = "Calibri"
        elif heading_font:
            _set_run_font(run, heading_font)
    if is_rtl:
        _set_rtl_paragraph(title_para)

    # Cover info — only show non-empty fields
    student = cover_data.get("student", "")
    university = cover_data.get("university", "")
    instructor = cover_data.get("instructor", "")
    date_val = cover_data.get("date", "")

    for field_val, field_label in [
        (student, None), (university, None),
        (instructor, "Instructor"), (date_val, "Date")
    ]:
        if field_val and field_val.strip():
            text = f"{field_label}: {field_val}" if field_label else field_val
            _add_styled_paragraph(doc, text, font_name=body_font, size_pt=body_size, is_rtl=is_rtl, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Always page break after cover
    doc.add_page_break()

    # === ABSTRACT ===
    abstract = report_data.get("abstract", "")
    if abstract:
        _add_styled_heading(doc, "Abstract", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=False)
        _add_styled_paragraph(doc, abstract, font_name=body_font, size_pt=body_size, is_rtl=is_rtl)
        doc.add_page_break()

    # === TABLE OF CONTENTS ===
    sections_data = report_data.get("sections", [])
    if sections_data:
        _add_styled_heading(doc, "Table of Contents", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)

        def strip_heading_label(text):
            if text.startswith("[Heading"):
                parts = text.split("]", 1)
                if len(parts) > 1:
                    return parts[1].strip()
            return text

        for section in sections_data:
            heading = strip_heading_label(section.get("heading", ""))
            _add_styled_paragraph(doc, heading, font_name=body_font, size_pt=body_size, is_rtl=is_rtl, style="List Number")

            for subsection in section.get("subsections", []):
                sub_heading = strip_heading_label(subsection.get("heading", ""))
                _add_styled_paragraph(doc, sub_heading, font_name=body_font, size_pt=body_size, is_rtl=is_rtl, style="List Number 2")

    # === MAIN CONTENT ===
    for i, section in enumerate(sections_data):
        heading = section.get("heading", "")
        content = section.get("content", "")

        # Each major section starts on a new page
        _add_styled_heading(doc, heading, level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)

        if content:
            _add_styled_paragraph(doc, content, font_name=body_font, size_pt=body_size, is_rtl=is_rtl)

        for subsection in section.get("subsections", []):
            sub_heading = subsection.get("heading", "")
            sub_content = subsection.get("content", "")

            _add_styled_heading(doc, sub_heading, level=2, font_name=heading_font, is_rtl=is_rtl, page_break=False)

            if sub_content:
                _add_styled_paragraph(doc, sub_content, font_name=body_font, size_pt=body_size, is_rtl=is_rtl)

            for subsubsection in subsection.get("subsubsections", []):
                subsub_heading = subsubsection.get("heading", "")
                subsub_content = subsubsection.get("content", "")

                _add_styled_heading(doc, subsub_heading, level=3, font_name=heading_font, is_rtl=is_rtl, page_break=False)

                if subsub_content:
                    _add_styled_paragraph(doc, subsub_content, font_name=body_font, size_pt=body_size, is_rtl=is_rtl)

    # === CONCLUSION ===
    conclusion = report_data.get("conclusion", "")
    if conclusion:
        _add_styled_heading(doc, "Conclusion", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)
        _add_styled_paragraph(doc, conclusion, font_name=body_font, size_pt=body_size, is_rtl=is_rtl)

    # === REFERENCES ===
    references = report_data.get("references", [])
    if references:
        _add_styled_heading(doc, "References", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)
        for ref in references:
            _add_styled_paragraph(doc, ref, font_name=body_font, size_pt=body_size, is_rtl=is_rtl, style="List Bullet")

    # === TABLES ===
    tables = report_data.get("tables", [])
    if tables:
        _add_styled_heading(doc, "Data Tables", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)
        for table_data in tables:
            if table_data.get("title"):
                _add_styled_heading(doc, table_data["title"], level=2, font_name=heading_font, is_rtl=is_rtl)
            add_table_to_doc(doc, table_data, font_name=body_font, is_rtl=is_rtl)

    # === SMART STRUCTURES ===
    smart_structures = report_data.get("smart_structures", [])
    if smart_structures:
        _add_styled_heading(doc, "Structured Analysis", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)
        for structure in smart_structures:
            add_smart_structure_to_doc(doc, structure, font_name=body_font, is_rtl=is_rtl)

    # === CHARTS ===
    charts = report_data.get("charts", [])
    if charts:
        _add_styled_heading(doc, "Data Analysis", level=1, font_name=heading_font, is_rtl=is_rtl, page_break=True)
        for chart in charts:
            add_chart_description_to_doc(doc, chart, font_name=body_font, is_rtl=is_rtl)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
